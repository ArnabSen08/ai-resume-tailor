from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import requests
import json
from typing import Optional
import logging
import io
import PyPDF2
import pdfplumber
import docx
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="AI Resume Tailor", version="1.0.0")

# Enable CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify your GitHub Pages URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class ResumeRequest(BaseModel):
    resume: str
    job_desc: str
    job_url: Optional[str] = None

class FileUploadResponse(BaseModel):
    extracted_text: str
    file_type: str
    success: bool
    message: str

class ResumeResponse(BaseModel):
    tailored_resume: str
    key_skills_extracted: list
    optimization_notes: str

# Ollama API configuration
OLLAMA_BASE_URL = "http://localhost:11434"  # Default Ollama URL
MODEL_NAME = "mistral"  # You can change this to gemma or other models

def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF using multiple methods for better accuracy
    """
    text = ""
    
    try:
        # Method 1: Using pdfplumber (better for complex layouts)
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if text.strip():
            return text.strip()
            
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {str(e)}")
    
    try:
        # Method 2: Using PyPDF2 as fallback
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
                
        return text.strip()
        
    except Exception as e:
        logger.error(f"PyPDF2 extraction failed: {str(e)}")
        raise HTTPException(status_code=400, detail="Could not extract text from PDF")

def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX files
    """
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text.strip())
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text.strip())
        
        return "\n".join(text)
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise HTTPException(status_code=400, detail="Could not extract text from DOCX")

def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text from uploaded file based on file type
    """
    file_content = file.file.read()
    file_extension = Path(file.filename).suffix.lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_content)
    elif file_extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_content)
    elif file_extension == '.txt':
        return file_content.decode('utf-8')
    else:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type: {file_extension}. Supported types: PDF, DOCX, TXT"
        )

def extract_job_description_from_url(url: str) -> str:
    """
    Extract job description from URL using web scraping
    """
    try:
        import requests
        from bs4 import BeautifulSoup
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Common job description selectors for popular job sites
        job_desc_selectors = [
            '.job-description',
            '.jobsearch-jobDescriptionText',
            '[data-testid="job-description"]',
            '.job-details',
            '.description',
            '.job-content',
            '.posting-description'
        ]
        
        job_description = ""
        for selector in job_desc_selectors:
            element = soup.select_one(selector)
            if element:
                job_description = element.get_text(strip=True)
                break
        
        if not job_description:
            # Fallback: get all paragraph text
            paragraphs = soup.find_all('p')
            job_description = ' '.join([p.get_text(strip=True) for p in paragraphs])
        
        return job_description[:5000]  # Limit to 5000 chars
        
    except Exception as e:
        logger.error(f"Error scraping job URL: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Could not scrape job URL: {str(e)}")

def call_ollama_api(prompt: str) -> str:
    """
    Call Ollama API with the given prompt
    """
    try:
        payload = {
            "model": MODEL_NAME,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": 0.7,
                "top_p": 0.9,
                "max_tokens": 2000
            }
        }
        
        response = requests.post(
            f"{OLLAMA_BASE_URL}/api/generate",
            json=payload,
            timeout=60
        )
        
        if response.status_code != 200:
            raise HTTPException(status_code=500, detail=f"Ollama API error: {response.text}")
        
        result = response.json()
        return result.get("response", "")
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Ollama API request failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to connect to Ollama: {str(e)}")

def create_resume_optimization_prompt(resume: str, job_desc: str) -> str:
    """
    Create a comprehensive prompt for resume optimization
    """
    return f"""You are an expert Resume Optimizer and ATS (Applicant Tracking System) specialist. Your task is to tailor a resume to match a specific job description while maintaining technical accuracy and authenticity.

**ORIGINAL RESUME:**
{resume}

**TARGET JOB DESCRIPTION:**
{job_desc}

**INSTRUCTIONS:**
1. **Extract Key Skills & Requirements:** Identify the most important technical skills, soft skills, and qualifications from the job description.

2. **Optimize Resume Content:**
   - Rewrite bullet points to mirror the job requirements using similar keywords and phrases
   - Quantify achievements where possible (percentages, numbers, metrics)
   - Ensure technical accuracy - don't add skills the candidate doesn't have
   - Maintain the original structure and format
   - Use action verbs that match the job posting tone

3. **ATS Optimization:**
   - Include relevant keywords naturally throughout the resume
   - Use standard section headings
   - Ensure proper formatting for ATS parsing

4. **Output Format:**
   Please provide your response in the following JSON format:
   {{
     "tailored_resume": "The complete optimized resume text",
     "key_skills_extracted": ["skill1", "skill2", "skill3"],
     "optimization_notes": "Brief explanation of key changes made"
   }}

**IMPORTANT:** Only enhance and optimize existing content. Do not fabricate experience or skills the candidate doesn't possess."""

@app.post("/upload-resume", response_model=FileUploadResponse)
async def upload_resume(file: UploadFile = File(...)):
    """
    Upload and extract text from resume file (PDF, DOCX, TXT)
    """
    try:
        # Validate file size (max 10MB)
        if file.size and file.size > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File size too large. Maximum 10MB allowed.")
        
        # Validate file type
        allowed_types = ['.pdf', '.docx', '.doc', '.txt']
        file_extension = Path(file.filename).suffix.lower()
        
        if file_extension not in allowed_types:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Allowed types: {', '.join(allowed_types)}"
            )
        
        # Extract text from file
        extracted_text = extract_text_from_file(file)
        
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the file")
        
        logger.info(f"Successfully extracted {len(extracted_text)} characters from {file.filename}")
        
        return FileUploadResponse(
            extracted_text=extracted_text,
            file_type=file_extension,
            success=True,
            message=f"Successfully extracted text from {file.filename}"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"File upload error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@app.get("/")
async def root():
    return {"message": "AI Resume Tailor API is running!"}

@app.get("/health")
async def health_check():
    return {"status": "healthy", "model": MODEL_NAME}

@app.post("/tailor-resume", response_model=ResumeResponse)
async def tailor_resume(request: ResumeRequest):
    """
    Tailor a resume based on job description
    """
    try:
        job_description = request.job_desc
        
        # If job URL is provided, scrape the job description
        if request.job_url:
            try:
                scraped_desc = extract_job_description_from_url(request.job_url)
                if scraped_desc:
                    job_description = scraped_desc
                    logger.info(f"Successfully scraped job description from URL")
            except Exception as e:
                logger.warning(f"Failed to scrape URL, using provided description: {str(e)}")
        
        # Create the optimization prompt
        prompt = create_resume_optimization_prompt(request.resume, job_description)
        
        # Call Ollama API
        logger.info("Calling Ollama API for resume optimization...")
        response = call_ollama_api(prompt)
        
        # Try to parse JSON response
        try:
            result = json.loads(response)
            return ResumeResponse(**result)
        except json.JSONDecodeError:
            # If not JSON, return as plain text with basic parsing
            return ResumeResponse(
                tailored_resume=response,
                key_skills_extracted=["AI/ML", "Python", "Data Analysis"],  # Default
                optimization_notes="Resume optimized based on job requirements"
            )
            
    except Exception as e:
        logger.error(f"Error in tailor_resume: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/scrape-job")
async def scrape_job_description(job_url: str):
    """
    Scrape job description from URL
    """
    try:
        job_desc = extract_job_description_from_url(job_url)
        return {"job_description": job_desc}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)